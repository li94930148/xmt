from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports" / "output" / "leadership_report_v1"
ASSET_DIR = REPORT_DIR / "assets"
DOCX_PATH = REPORT_DIR / "岚曜新媒体平台项目领导汇报初稿.docx"
HTML_PATH = REPORT_DIR / "岚曜新媒体平台项目领导汇报初稿.html"
META_PATH = REPORT_DIR / "report_meta.json"
LOGO_PATH = ROOT / "public" / "logo.png"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD_PATH = Path(r"C:\Windows\Fonts\msyhbd.ttc")

PRIMARY = (36, 99, 235)
PRIMARY_DARK = (22, 61, 168)
SECONDARY = (9, 122, 105)
TEXT = (31, 41, 55)
SUBTLE = (100, 116, 139)
BG = (247, 250, 252)
CARD = (255, 255, 255)
LINE = (226, 232, 240)
ACCENT = (245, 158, 11)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = FONT_BOLD_PATH if bold and FONT_BOLD_PATH.exists() else FONT_PATH
    return ImageFont.truetype(str(font_path), size=size)


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def count_files(path: Path, pattern: str = "*") -> int:
    return len([item for item in path.glob(pattern) if item.is_file()])


def count_recursive_files(path: Path, pattern: str = "*") -> int:
    return len([item for item in path.rglob(pattern) if item.is_file()])


def get_metrics() -> dict[str, int]:
    return {
        "route_modules": count_files(ROOT / "api" / "routes", "*.ts"),
        "page_modules": count_files(ROOT / "src" / "pages", "*.tsx"),
        "component_modules": count_recursive_files(ROOT / "src" / "components", "*.tsx"),
        "hook_modules": count_files(ROOT / "src" / "hooks", "*.ts"),
        "api_clients": count_files(ROOT / "src" / "api", "*.ts"),
    }


def create_canvas(size: tuple[int, int]) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, BG)
    return image, ImageDraw.Draw(image)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: tuple[int, int, int],
    outline: tuple[int, int, int] | None = None,
    radius: int = 24,
    width: int = 2,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def add_wrapped_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    font: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int] = TEXT,
    line_gap: int = 8,
    align: str = "left",
) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1
    lines: list[str] = []
    for raw_line in text.splitlines():
        if not raw_line:
            lines.append("")
            continue
        current = ""
        for char in raw_line:
            candidate = current + char
            width = draw.textbbox((0, 0), candidate, font=font)[2]
            if width <= max_width or not current:
                current = candidate
            else:
                lines.append(current)
                current = char
        if current:
            lines.append(current)

    line_height = draw.textbbox((0, 0), "示例", font=font)[3]
    total_height = len(lines) * line_height + max(len(lines) - 1, 0) * line_gap
    y = y1 + max(0, (y2 - y1 - total_height) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_width = bbox[2] - bbox[0]
        if align == "center":
            x = x1 + max(0, (max_width - line_width) // 2)
        elif align == "right":
            x = x2 - line_width
        else:
            x = x1
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height + line_gap


def draw_title_block(draw: ImageDraw.ImageDraw, title: str, subtitle: str, size: tuple[int, int]) -> None:
    width, _ = size
    draw.rectangle((0, 0, width, 110), fill=PRIMARY)
    draw.text((70, 28), title, font=load_font(34, bold=True), fill=(255, 255, 255))
    draw.text((70, 72), subtitle, font=load_font(18), fill=(230, 236, 255))


def draw_architecture(asset_path: Path) -> None:
    image, draw = create_canvas((1600, 900))
    draw_title_block(draw, "系统整体结构图", "从使用层、协同层到业务层与数据层，形成一套完整闭环平台", image.size)

    layer_titles = [
        ("使用层", "团队成员日常可见的页面与入口", PRIMARY),
        ("协同与控制层", "权限、审批、实时通知与消息机制", (14, 116, 144)),
        ("业务执行层", "围绕内容生产全过程的核心业务模块", (8, 126, 88)),
        ("数据与基础层", "本地部署、数据存储、导出、备份与服务支撑", (109, 40, 217)),
    ]
    y_positions = [170, 330, 500, 690]
    box_text = [
        ["首页看板", "选题管理", "创作管理", "拍摄 / 成片", "发布管理", "数据复盘"],
        ["角色权限", "审批流设计", "站内消息", "Socket 实时同步", "操作日志", "通知设置"],
        ["灵感池", "日历排期", "看板协作", "资源库", "成就体系", "抖音数据采集"],
        ["Express + SQLite", "JWT 认证", "导出与备份", "统一 API", "局域网/单机部署", "后续可扩展为通用产品"],
    ]

    for idx, ((title, subtitle, color), y) in enumerate(zip(layer_titles, y_positions)):
        rounded_box(draw, (60, y, 1540, y + 120), CARD, outline=LINE, radius=28, width=2)
        draw.rounded_rectangle((78, y + 18, 300, y + 100), radius=24, fill=color)
        add_wrapped_text(draw, title, (100, y + 28, 280, y + 62), load_font(24, bold=True), fill=(255, 255, 255))
        add_wrapped_text(draw, subtitle, (100, y + 58, 280, y + 94), load_font(14), fill=(230, 240, 255))

        start_x = 340
        for label in box_text[idx]:
            left = start_x
            right = left + 172
            rounded_box(draw, (left, y + 22, right, y + 98), (248, 250, 252), outline=color, radius=20, width=2)
            add_wrapped_text(draw, label, (left + 14, y + 34, right - 14, y + 88), load_font(18), align="center")
            start_x += 195

    image.save(asset_path)


def draw_workflow(asset_path: Path) -> None:
    image, draw = create_canvas((1600, 760))
    draw_title_block(draw, "项目运作流程图", "把选题到复盘串成一条线，减少信息断层和责任不清", image.size)

    steps = [
        ("1", "选题策划", "沉淀选题来源、方向和初步判断"),
        ("2", "审核通过", "统一把关，避免资源投入到低价值事项"),
        ("3", "创作生产", "脚本、大纲、素材、批注集中管理"),
        ("4", "拍摄执行", "按计划推进，任务与状态可追踪"),
        ("5", "成片处理", "围绕发布前准备进行衔接"),
        ("6", "发布管理", "统一记录发布动作，避免遗漏与重复"),
        ("7", "数据复盘", "把结果回连到选题和执行过程"),
    ]

    x = 70
    y = 250
    width = 190
    height = 260
    gap = 26
    for idx, (num, title, desc) in enumerate(steps):
        left = x + idx * (width + gap)
        right = left + width
        box_fill = (236, 245, 255) if idx % 2 == 0 else (239, 250, 246)
        border = PRIMARY if idx % 2 == 0 else SECONDARY
        rounded_box(draw, (left, y, right, y + height), box_fill, outline=border, radius=24, width=3)
        draw.ellipse((left + 62, y - 28, left + 128, y + 38), fill=border)
        add_wrapped_text(draw, num, (left + 78, y - 16, left + 112, y + 20), load_font(24, bold=True), fill=(255, 255, 255), align="center")
        add_wrapped_text(draw, title, (left + 18, y + 34, right - 18, y + 86), load_font(24, bold=True), align="center")
        add_wrapped_text(draw, desc, (left + 16, y + 96, right - 16, y + 220), load_font(16), line_gap=10)
        if idx < len(steps) - 1:
            arrow_y = y + height // 2
            arrow_x1 = right + 8
            arrow_x2 = right + gap - 10
            draw.line((arrow_x1, arrow_y, arrow_x2, arrow_y), fill=SUBTLE, width=6)
            draw.polygon([(arrow_x2, arrow_y), (arrow_x2 - 18, arrow_y - 12), (arrow_x2 - 18, arrow_y + 12)], fill=SUBTLE)

    draw.text((70, 610), "关键管理价值：每一步都对应明确状态、责任、记录与下一步触发条件。", font=load_font(22, bold=True), fill=PRIMARY_DARK)
    image.save(asset_path)


def draw_scope_chart(asset_path: Path, metrics: dict[str, int]) -> None:
    image, draw = create_canvas((1600, 900))
    draw_title_block(draw, "当前系统建设范围（代码结构事实）", "以下数量仅用于说明系统覆盖广度，不代表业务价值大小", image.size)

    items = [
        ("前端页面模块", metrics["page_modules"], PRIMARY),
        ("后端业务路由", metrics["route_modules"], (14, 116, 144)),
        ("通用组件", metrics["component_modules"], (8, 126, 88)),
        ("前端 API 封装", metrics["api_clients"], (109, 40, 217)),
        ("自定义 Hook", metrics["hook_modules"], (217, 70, 239)),
    ]
    max_value = max(value for _, value, _ in items)
    left = 220
    top = 210
    bar_max = 1100
    row_gap = 110

    for idx, (label, value, color) in enumerate(items):
        y = top + idx * row_gap
        draw.text((70, y + 20), label, font=load_font(24, bold=True), fill=TEXT)
        draw.rounded_rectangle((left, y, left + bar_max, y + 54), radius=20, fill=(232, 239, 247))
        bar_width = int(bar_max * value / max_value)
        draw.rounded_rectangle((left, y, left + bar_width, y + 54), radius=20, fill=color)
        draw.text((left + bar_max + 30, y + 10), str(value), font=load_font(28, bold=True), fill=color)

    notes = [
        "说明 1：系统不是单一页面，而是已经具备前后端、权限、流程、通知、数据等平台能力。",
        "说明 2：这张图更适合作为“建设完整度”的直观辅助，而不是强调代码量。",
    ]
    note_y = 760
    for note in notes:
        draw.text((72, note_y), note, font=load_font(20), fill=SUBTLE)
        note_y += 34
    image.save(asset_path)


def draw_roadmap(asset_path: Path) -> None:
    image, draw = create_canvas((1600, 860))
    draw_title_block(draw, "未来愿景与产品化路径", "从内部提效工具，逐步演进为可复制、可交付、可盈利的通用系统", image.size)

    phases = [
        ("阶段一：内部深耕", "围绕现有团队继续打磨", ["流程更稳定", "角色更清晰", "形成标准作业方式"], PRIMARY),
        ("阶段二：通用化改造", "把团队经验抽象成可配置能力", ["模块参数化", "流程模板化", "权限可配置"], (14, 116, 144)),
        ("阶段三：公司内复制", "服务更多团队或业务线", ["降低复制成本", "缩短搭建周期", "统一管理视角"], SECONDARY),
        ("阶段四：整体打包出售", "服务需要多环节协同的新媒体公司", ["整包交付", "私有化部署", "实施与运维服务"], (109, 40, 217)),
    ]

    left = 80
    top = 260
    width = 330
    height = 320
    gap = 36
    for idx, (title, subtitle, bullets, color) in enumerate(phases):
        x1 = left + idx * (width + gap)
        x2 = x1 + width
        rounded_box(draw, (x1, top, x2, top + height), CARD, outline=color, radius=28, width=3)
        draw.rectangle((x1, top, x2, top + 70), fill=color)
        add_wrapped_text(draw, title, (x1 + 18, top + 16, x2 - 18, top + 46), load_font(24, bold=True), fill=(255, 255, 255), align="center")
        add_wrapped_text(draw, subtitle, (x1 + 22, top + 88, x2 - 22, top + 138), load_font(18, bold=True), fill=TEXT, align="center")
        bullet_y = top + 162
        for bullet in bullets:
            draw.ellipse((x1 + 24, bullet_y + 10, x1 + 38, bullet_y + 24), fill=color)
            add_wrapped_text(draw, bullet, (x1 + 50, bullet_y, x2 - 24, bullet_y + 44), load_font(18), line_gap=4)
            bullet_y += 52
        if idx < len(phases) - 1:
            ay = top + height // 2
            ax1 = x2 + 8
            ax2 = x2 + gap - 10
            draw.line((ax1, ay, ax2, ay), fill=ACCENT, width=7)
            draw.polygon([(ax2, ay), (ax2 - 18, ay - 12), (ax2 - 18, ay + 12)], fill=ACCENT)

    draw.text((80, 660), "当前最值得继续投入的方向：优化整体系统架构、整合各模块，并统一系统风格。", font=load_font(22, bold=True), fill=PRIMARY_DARK)
    image.save(asset_path)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for style_name, size, bold, color in [
        ("Title", 24, True, "1D4ED8"),
        ("Heading 1", 16, True, "1E3A8A"),
        ("Heading 2", 13, True, "0F766E"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text: str, *, bold: bool = False, size: int | None = None, color: str | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(6)
        add_run(p, "• ", bold=True, color="1D4ED8")
        add_run(p, item)


def add_image(doc: Document, path: Path, width: float = 6.7) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def add_summary_table(doc: Document) -> None:
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目名称", "岚曜新媒体平台"),
        ("项目定位", "一套服务新媒体内容团队的全流程协作管理系统，覆盖选题、审核、创作、拍摄、发布、复盘。"),
        ("当前使用角色", "文案、摄影、剪辑已经在实际使用。"),
        ("阶段成果", "“泰安的十大名人墓葬”选题已在平台内完成全流程跑通。"),
        ("已体现价值", "减少约 80% 沟通成本，减少整理思路、重复提交选题文档、重复修改等动作。"),
        ("未来方向", "继续优化整体架构并统一风格，面向需要多环节协同的新媒体公司整体打包出售。"),
    ]
    for idx, (left, right) in enumerate(rows):
        table.cell(idx, 0).text = left
        table.cell(idx, 1).text = right
        table.cell(idx, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(idx, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(table.cell(idx, 0), "EAF2FF")


def add_advantage_table(doc: Document) -> None:
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["维度", "当前优势", "为什么难以替代"]
    for col, text in enumerate(headers):
        table.cell(0, col).text = text
        set_cell_shading(table.cell(0, col), "DCEBFF")
    rows = [
        ("流程闭环", "从选题到复盘贯穿在同一系统内", "不是单点工具，而是把前后环节接起来，替代零散协作方式"),
        ("管理可视化", "每个状态、责任和进度都有记录", "领导和负责人可以随时看到推进情况，减少口头追踪"),
        ("经验沉淀", "脚本、灵感、资源、审批方式都能留下来", "团队能力不再只依赖个人记忆，便于复制和交接"),
        ("协同效率", "消息、通知、权限、日历、看板形成联动", "减少跨群确认、重复问进度、版本不一致等管理成本"),
        ("数据回连", "复盘数据能回到前面的选题与执行过程", "不仅看结果，也能分析为什么做成或没做成"),
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            table.cell(row_idx, col_idx).text = text
            table.cell(row_idx, col_idx).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(table.cell(row_idx, 0), "F8FAFC")


def add_pending_info_table(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers = ["已确认信息", "当前内容", "汇报用途"]
    for col, text in enumerate(headers):
        table.cell(0, col).text = text
        set_cell_shading(table.cell(0, col), "FEF3C7")
    rows = [
        ("落地范围", "文案、摄影、剪辑已在使用", "证明项目已进入实际落地阶段"),
        ("效率改善", "减少约 80% 沟通成本", "让领导直观看到投入产出比"),
        ("典型案例", "“泰安的十大名人墓葬”已全流程跑通", "证明系统可以支撑真实业务闭环"),
        ("当前重点", "优化整体系统架构、整合模块、统一风格", "为下一阶段资源投入提供方向"),
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            table.cell(row_idx, col_idx).text = text


def build_doc(metrics: dict[str, int]) -> None:
    doc = Document()
    set_document_defaults(doc)

    doc.add_heading("一、项目摘要", level=1)
    p = doc.add_paragraph()
    add_run(p, "一句话定义：", bold=True, color="1D4ED8")
    add_run(p, "岚曜新媒体平台是一套面向新媒体内容团队的全流程协作管理系统，用来把选题、审核、创作、拍摄、发布和复盘放进同一条业务链路。")
    add_summary_table(doc)

    doc.add_heading("二、为什么要做这个项目", level=1)
    add_bullets(
        doc,
        [
            "内容生产链路长、协作角色多，如果没有统一系统，信息很容易分散在聊天记录、表格和口头沟通中。",
            "选题、脚本、执行、发布、复盘由不同角色接力完成，一旦缺少统一状态和责任归属，就会出现遗漏、重复确认和推进不透明。",
            "很多团队最缺的不是单点工具，而是一套能把“计划 - 执行 - 结果”连起来的管理底座。",
            "当前实际使用反馈已经说明，这套系统能明显减少反复沟通、整理思路、重复提交选题文档以及来回修改等低效率动作。",
        ],
    )

    doc.add_heading("三、这个项目整体是什么", level=1)
    add_bullets(
        doc,
        [
            "它不是单一页面，也不是单一表单，而是一套完整的前后端系统。",
            "前端负责日常操作入口和协同界面，后端负责业务规则、权限、消息、数据存储与接口能力。",
            "系统当前已具备内容流程模块、效率工具模块、数据管理模块和平台支撑能力。",
        ],
    )
    add_image(doc, ASSET_DIR / "architecture.png")

    doc.add_heading("四、当前系统建设范围", level=1)
    p = doc.add_paragraph()
    add_run(p, "基于当前代码库可直接确认的事实：", bold=True)
    add_run(
        p,
        f" 已具备 {metrics['page_modules']} 个前端页面模块、{metrics['route_modules']} 个后端业务路由、"
        f"{metrics['component_modules']} 个通用组件、{metrics['api_clients']} 个前端 API 封装、"
        f"{metrics['hook_modules']} 个自定义 Hook。",
    )
    add_image(doc, ASSET_DIR / "scope_chart.png")

    doc.add_heading("五、项目是怎么运作的", level=1)
    add_bullets(
        doc,
        [
            "系统主线以“选题”为起点，通过审核后进入后续创作、执行、发布和复盘环节。",
            "各环节之间不是孤立跳转，而是通过统一状态、数据记录和权限机制实现接续。",
            "团队不需要反复手工同步“进行到哪一步了”，系统本身就承担了流程连接器的角色。",
            "“泰安的十大名人墓葬”这一选题已经验证了平台具备支撑完整业务闭环的能力。",
        ],
    )
    add_image(doc, ASSET_DIR / "workflow.png")

    doc.add_heading("六、这个项目的核心优点", level=1)
    add_bullets(
        doc,
        [
            "全链路闭环：把内容生产全过程放在同一套系统里，减少中间断层。",
            "标准化管理：通过状态流转、审批设计、权限控制，把经验变成制度化动作。",
            "实时协同：消息、通知、日志、实时同步机制增强团队协作效率。",
            "数据沉淀：系统不只记录结果，还能把前置过程沉淀下来，为复盘和优化提供依据。",
            "可控部署：当前采用 Express + SQLite 的轻量部署方式，适合内部快速落地和持续迭代。",
            "实际价值已初步体现：目前反馈可减少约 80% 沟通成本，并减少整理思路、反复提报和重复修改等动作。",
        ],
    )

    doc.add_heading("七、对团队和公司的不可替代优势", level=1)
    add_advantage_table(doc)

    doc.add_heading("八、未来愿景：从内部系统走向通用产品", level=1)
    add_bullets(
        doc,
        [
            "如果继续沿着“可配置、可复制、可部署”的方向改造，这套系统有机会从团队专用工具演进成通用系统。",
            "通用化的关键，不是简单增加功能，而是把当前团队经验抽象成模板化、参数化、角色化能力。",
            "当前最值得继续投入的方向，是进一步优化整体系统架构、整合各模块，并统一系统风格。",
            "更适合的目标客户，是新媒体等需要多环节协同合作的公司，未来可考虑以整体打包方案出售。",
        ],
    )
    add_image(doc, ASSET_DIR / "roadmap.png")

    doc.add_heading("九、给领导汇报时建议重点强调的结论", level=1)
    add_bullets(
        doc,
        [
            "这不是一个零散工具集合，而是一套已经具备平台雏形的业务系统。",
            "它最大的价值不是“多了一个系统”，而是把团队经验、流程规则和协同方式沉淀下来。",
            "对团队来说，它提升的是执行效率和管理透明度；对公司来说，它沉淀的是未来可复制、可扩展、可盈利的数字化能力。",
            "从当前使用反馈看，这个项目已经有继续打磨和扩大应用范围的基础。",
        ],
    )

    doc.add_heading("十、当前已确认的关键事实", level=1)
    p = doc.add_paragraph()
    add_run(p, "以下内容已经可以作为正式汇报口径使用：", bold=True, color="92400E")
    add_pending_info_table(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理", size=9, color="64748B")
    doc.save(DOCX_PATH)


def build_assets(metrics: dict[str, int]) -> None:
    draw_architecture(ASSET_DIR / "architecture.png")
    draw_workflow(ASSET_DIR / "workflow.png")
    draw_scope_chart(ASSET_DIR / "scope_chart.png", metrics)
    draw_roadmap(ASSET_DIR / "roadmap.png")


def save_meta(metrics: dict[str, int]) -> None:
    meta = {
        "project": "岚曜新媒体平台",
        "metrics": metrics,
        "docx": str(DOCX_PATH),
        "html": str(HTML_PATH),
        "assets": [str(path) for path in sorted(ASSET_DIR.glob("*.png"))],
        "notes": [
            "本稿仅使用代码库、架构文档和用户已确认业务信息。",
            "如需正式版，可继续补充更多案例、截图和领导关注指标。",
        ],
    }
    META_PATH.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")


def build_html(metrics: dict[str, int]) -> None:
    assets = {
        "architecture": (ASSET_DIR / "architecture.png").as_uri(),
        "workflow": (ASSET_DIR / "workflow.png").as_uri(),
        "scope": (ASSET_DIR / "scope_chart.png").as_uri(),
        "roadmap": (ASSET_DIR / "roadmap.png").as_uri(),
    }
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>岚曜新媒体平台项目领导汇报初稿</title>
  <style>
    @page {{ size: A4; margin: 14mm; }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; font-family: "Microsoft YaHei", "Segoe UI", sans-serif; background: #f3f7fb; color: #1f2937; }}
    .page {{ width: 210mm; min-height: 297mm; margin: 0 auto 16px; background: white; padding: 18mm 16mm; position: relative; overflow: hidden; }}
    .page:last-child {{ margin-bottom: 0; }}
    .top-band {{ margin: -18mm -16mm 14px; padding: 16px 16mm 12px; background: linear-gradient(90deg, #2563eb 0%, #1d4ed8 55%, #0f766e 100%); color: white; }}
    .top-band h2 {{ margin: 0 0 6px; font-size: 24px; }}
    .top-band p {{ margin: 0; font-size: 12px; color: rgba(255,255,255,0.88); }}
    .section-title {{ margin: 0 0 10px; font-size: 18px; color: #1e3a8a; }}
    .lead {{ font-size: 14px; line-height: 1.85; color: #334155; margin: 0 0 10px; }}
    .summary-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 12px; }}
    .card {{ border: 1px solid #dbeafe; border-radius: 16px; padding: 12px 14px; background: #f8fbff; }}
    .card h4 {{ margin: 0 0 6px; font-size: 13px; color: #1d4ed8; }}
    .card p {{ margin: 0; font-size: 12px; line-height: 1.75; color: #334155; }}
    ul {{ margin: 8px 0 0 18px; padding: 0; line-height: 1.8; color: #334155; font-size: 13px; }}
    li {{ margin-bottom: 4px; }}
    .image {{ width: 100%; border: 1px solid #e2e8f0; border-radius: 16px; margin: 12px 0 0; overflow: hidden; background: #f8fafc; }}
    .image img {{ width: 100%; display: block; }}
    .metric-row {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 10px; margin: 12px 0; }}
    .metric {{ background: linear-gradient(180deg, #f8fbff 0%, #eef6ff 100%); border: 1px solid #dbeafe; border-radius: 16px; padding: 12px 10px; text-align: center; }}
    .metric strong {{ display: block; font-size: 24px; color: #1d4ed8; margin-bottom: 4px; }}
    .metric span {{ font-size: 11px; color: #475569; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 12px; line-height: 1.7; }}
    th, td {{ border: 1px solid #dbe3ef; padding: 9px 10px; vertical-align: top; text-align: left; }}
    th {{ background: #eaf2ff; color: #1e3a8a; }}
    .warn th {{ background: #fef3c7; color: #92400e; }}
    .footer-note {{ position: absolute; left: 16mm; right: 16mm; bottom: 10mm; text-align: center; color: #64748b; font-size: 10px; }}
  </style>
</head>
<body>
  <section class="page">
    <div class="top-band">
      <h2>一、项目摘要</h2>
      <p>先让领导在最短时间内知道：这是什么、为什么值得继续投入</p>
    </div>
    <p class="lead"><strong>一句话定义：</strong>岚曜新媒体平台是一套面向新媒体内容团队的全流程协作管理系统，用来把选题、审核、创作、拍摄、发布和复盘放进同一条业务链路。</p>
    <div class="summary-grid">
      <div class="card"><h4>项目名称</h4><p>岚曜新媒体平台</p></div>
      <div class="card"><h4>项目定位</h4><p>不是单一表单或单一页面，而是一套贯穿内容生产全过程的协同系统。</p></div>
      <div class="card"><h4>当前使用角色</h4><p>文案、摄影、剪辑已经在实际使用。</p></div>
      <div class="card"><h4>阶段成果</h4><p>“泰安的十大名人墓葬”选题已在平台内完成全流程跑通。</p></div>
      <div class="card"><h4>已体现价值</h4><p>减少约 80% 沟通成本，减少整理思路、重复提交选题文档、重复修改等动作。</p></div>
      <div class="card"><h4>未来方向</h4><p>继续优化整体架构并统一风格，面向需要多环节协同的新媒体公司整体打包出售。</p></div>
    </div>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>

  <section class="page">
    <div class="top-band">
      <h2>二、为什么要做这个项目</h2>
      <p>从管理痛点切入，比直接讲技术更容易让领导理解项目必要性</p>
    </div>
    <ul>
      <li>内容生产链路长、协作角色多，如果没有统一系统，信息很容易分散在聊天记录、表格和口头沟通中。</li>
      <li>选题、脚本、执行、发布、复盘由不同角色接力完成，一旦缺少统一状态和责任归属，就会出现遗漏、重复确认和推进不透明。</li>
      <li>很多团队最缺的不是单点工具，而是一套能把“计划 - 执行 - 结果”连起来的管理底座。</li>
      <li>当前实际使用反馈已经说明，这套系统能明显减少反复沟通、整理思路、重复提交选题文档以及来回修改等低效率动作。</li>
    </ul>
    <h3 class="section-title">这个项目整体是什么</h3>
    <ul>
      <li>它不是单一页面，也不是单一表单，而是一套完整的前后端系统。</li>
      <li>前端负责日常操作入口和协同界面，后端负责业务规则、权限、消息、数据存储与接口能力。</li>
      <li>系统当前已具备内容流程模块、效率工具模块、数据管理模块和平台支撑能力。</li>
    </ul>
    <div class="image"><img src="{assets["architecture"]}" alt="系统整体结构图" /></div>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>

  <section class="page">
    <div class="top-band">
      <h2>三、当前系统建设范围</h2>
      <p>用事实说明：这已经不是一个概念原型，而是具备完整度的平台雏形</p>
    </div>
    <div class="metric-row">
      <div class="metric"><strong>{metrics["page_modules"]}</strong><span>前端页面模块</span></div>
      <div class="metric"><strong>{metrics["route_modules"]}</strong><span>后端业务路由</span></div>
      <div class="metric"><strong>{metrics["component_modules"]}</strong><span>通用组件</span></div>
      <div class="metric"><strong>{metrics["api_clients"]}</strong><span>前端 API 封装</span></div>
      <div class="metric"><strong>{metrics["hook_modules"]}</strong><span>自定义 Hook</span></div>
    </div>
    <p class="lead">以上数量仅用于说明建设覆盖范围，不代表业务价值本身。但它能反映出：当前系统已经形成页面层、接口层、组件层、能力层的基础平台结构。</p>
    <div class="image"><img src="{assets["scope"]}" alt="当前系统建设范围" /></div>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>

  <section class="page">
    <div class="top-band">
      <h2>四、项目是怎么运作的</h2>
      <p>把“流程闭环”讲明白，是领导判断项目价值的关键</p>
    </div>
    <ul>
      <li>系统主线以“选题”为起点，通过审核后进入后续创作、执行、发布和复盘环节。</li>
      <li>各环节之间不是孤立跳转，而是通过统一状态、数据记录和权限机制实现接续。</li>
      <li>团队不需要反复手工同步“进行到哪一步了”，系统本身就承担了流程连接器的角色。</li>
      <li>“泰安的十大名人墓葬”这一选题已经验证了平台具备支撑完整业务闭环的能力。</li>
    </ul>
    <div class="image"><img src="{assets["workflow"]}" alt="项目运作流程图" /></div>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>

  <section class="page">
    <div class="top-band">
      <h2>五、这个项目的核心优点</h2>
      <p>重点不在“功能多”，而在于“能不能形成组织能力”</p>
    </div>
    <ul>
      <li>全链路闭环：把内容生产全过程放在同一套系统里，减少中间断层。</li>
      <li>标准化管理：通过状态流转、审批设计、权限控制，把经验变成制度化动作。</li>
      <li>实时协同：消息、通知、日志、实时同步机制增强团队协作效率。</li>
      <li>数据沉淀：系统不只记录结果，还能把前置过程沉淀下来，为复盘和优化提供依据。</li>
      <li>可控部署：当前采用 Express + SQLite 的轻量部署方式，适合内部快速落地和持续迭代。</li>
      <li>实际价值已初步体现：目前反馈可减少约 80% 沟通成本，并减少整理思路、反复提报和重复修改等动作。</li>
    </ul>
    <h3 class="section-title">对团队和公司的不可替代优势</h3>
    <table>
      <thead>
        <tr><th>维度</th><th>当前优势</th><th>为什么难以替代</th></tr>
      </thead>
      <tbody>
        <tr><td>流程闭环</td><td>从选题到复盘贯穿在同一系统内</td><td>不是单点工具，而是把前后环节接起来，替代零散协作方式</td></tr>
        <tr><td>管理可视化</td><td>每个状态、责任和进度都有记录</td><td>领导和负责人可以随时看到推进情况，减少口头追踪</td></tr>
        <tr><td>经验沉淀</td><td>脚本、灵感、资源、审批方式都能留下来</td><td>团队能力不再只依赖个人记忆，便于复制和交接</td></tr>
        <tr><td>协同效率</td><td>消息、通知、权限、日历、看板形成联动</td><td>减少跨群确认、重复问进度、版本不一致等管理成本</td></tr>
        <tr><td>数据回连</td><td>复盘数据能回到前面的选题与执行过程</td><td>不仅看结果，也能分析为什么做成或没做成</td></tr>
      </tbody>
    </table>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>

  <section class="page">
    <div class="top-band">
      <h2>六、未来愿景：从内部系统走向通用产品</h2>
      <p>先说明这套系统为什么值得继续做，再说明未来为什么有商业化潜力</p>
    </div>
    <ul>
      <li>如果继续沿着“可配置、可复制、可部署”的方向改造，这套系统有机会从团队专用工具演进成通用系统。</li>
      <li>通用化的关键，不是简单增加功能，而是把当前团队经验抽象成模板化、参数化、角色化能力。</li>
      <li>当前最值得继续投入的方向，是进一步优化整体系统架构、整合各模块，并统一系统风格。</li>
      <li>更适合的目标客户，是新媒体等需要多环节协同合作的公司，未来可考虑以整体打包方案出售。</li>
    </ul>
    <div class="image"><img src="{assets["roadmap"]}" alt="未来愿景与产品化路径" /></div>
    <h3 class="section-title">给领导汇报时建议重点强调的结论</h3>
    <ul>
      <li>这不是一个零散工具集合，而是一套已经具备平台雏形的业务系统。</li>
      <li>它最大的价值不是“多了一个系统”，而是把团队经验、流程规则和协同方式沉淀下来。</li>
      <li>对团队来说，它提升的是执行效率和管理透明度；对公司来说，它沉淀的是未来可复制、可扩展、可盈利的数字化能力。</li>
      <li>从当前使用反馈看，这个项目已经有继续打磨和扩大应用范围的基础。</li>
    </ul>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>

  <section class="page">
    <div class="top-band">
      <h2>七、当前已确认的关键事实</h2>
      <p>这部分已经可以作为正式汇报口径使用</p>
    </div>
    <table class="warn">
      <thead>
        <tr><th>已确认信息</th><th>当前内容</th><th>汇报用途</th></tr>
      </thead>
      <tbody>
        <tr><td>落地范围</td><td>文案、摄影、剪辑已在使用</td><td>证明项目已进入实际落地阶段</td></tr>
        <tr><td>效率改善</td><td>减少约 80% 沟通成本</td><td>让领导直观看到投入产出比</td></tr>
        <tr><td>典型案例</td><td>“泰安的十大名人墓葬”已全流程跑通</td><td>证明系统可以支撑真实业务闭环</td></tr>
        <tr><td>当前重点</td><td>优化整体系统架构、整合模块、统一风格</td><td>为下一阶段资源投入提供方向</td></tr>
      </tbody>
    </table>
    <div class="footer-note">内部汇报初稿｜基于代码库、架构文档与已确认业务信息整理</div>
  </section>
</body>
</html>
"""
    HTML_PATH.write_text(html, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    metrics = get_metrics()
    build_assets(metrics)
    build_doc(metrics)
    build_html(metrics)
    save_meta(metrics)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
